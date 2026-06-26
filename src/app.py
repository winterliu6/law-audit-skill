"""
法务审核系统主应用入口
FastAPI + uvicorn，监听0.0.0.0:3330（仅局域网）
"""
import os
import asyncio
import json
import shutil
import secrets
import bcrypt
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Form, Request, Query
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session

from .config import BASE_DIR, WEB_DIR, REACT_WEB_DIR, USER_CONTRACT_DIR, KB_UPLOAD_DIR, SAVE_DIR, HOST, PORT
from .database import init_db, get_db, User, Contract, Consultation, WorkOrder, AuditRecord, DeviceBinding, Organization, GuestSession, ContractTemplate, GeneratedContract
from .user_auth.auth import (register, login, get_current_user, get_optional_user, require_admin,
                             set_auth_cookie, hash_fingerprint, create_guest_session)
from .llm_sync.sync import start_sync, get_llm_config, call_llm
from .expert_link.expert_bridge import get_expert_context
from .law_kb.knowledge_base import init_base_knowledge, search as kb_search, parse_pdf, parse_docx, add_document, rebuild_index
from .contract_engine.audit import parse_contract, analyze_contract, generate_report, extract_key_info
from .docx_comments import generate_annotated_docx
from .contract_template.api import router as contract_template_router
from .work_dispatch.dispatcher import create_order, get_orders, accept_order, complete_order, return_order, transfer_order

app = FastAPI(title="法务审核系统", version="2.1.0")

app.add_middleware(
    CORSMiddleware,
    # 生产环境加固：如对外暴露，建议将 allow_origins 从 ["*"] 改为具体域名列表
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files: React app assets + legacy web_front
if REACT_WEB_DIR.exists():
    app.mount("/assets", StaticFiles(directory=str(REACT_WEB_DIR / "assets")), name="react_assets")
app.mount("/static", StaticFiles(directory=str(WEB_DIR)), name="static")

app.include_router(contract_template_router)


# SPA helper: serve React index.html for all frontend pages
def _spa_response():
    index_path = REACT_WEB_DIR / "index.html"
    if index_path.exists():
        return FileResponse(str(index_path))
    return HTMLResponse("<h1>524d7aef672a67845efaFf0c8bf78fd0884c cd web_app && npm run build</h1>", status_code=503)
@app.on_event("startup")
async def startup():
    init_db()
    init_base_knowledge()
    start_sync()
    print(f"[启动] 法务审核系统 v2.0 就绪: http://{HOST}:{PORT}")


# ==================== 页面路由 ====================

@app.get("/", response_class=HTMLResponse)
async def root():
    return _spa_response()
@app.get("/audit", response_class=HTMLResponse)
async def audit_page():
    return _spa_response()
@app.get("/work-order", response_class=HTMLResponse)
async def work_order_page():
    return _spa_response()
@app.get("/history", response_class=HTMLResponse)
async def history_page():
    return _spa_response()
@app.get("/contract-template", response_class=HTMLResponse)
async def contract_template_page():
    return _spa_response()
@app.get("/admin", response_class=HTMLResponse)
async def admin_page():
    return _spa_response()

@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard_page():
    return _spa_response()

# ==================== 认证接口 ====================

@app.post("/api/register")
async def api_register(
    username: str = Form(...), password: str = Form(...),
    company: str = Form(""), department: str = Form(""), full_name: str = Form(""),
    db: Session = Depends(get_db)
):
    result = register(username, password, db, company=company, department=department, full_name=full_name)
    return {"code": 0, "msg": "注册成功", "data": result}


@app.post("/api/login")
async def api_login(request: Request, username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    body = await request.form()
    fp = body.get("device_fingerprint", "unknown")
    result = login(username, password, fp, db)
    response = JSONResponse({"code": 0, "msg": "登录成功", "data": result["user"]})
    set_auth_cookie(response, result["token"])
    return response


@app.get("/api/me")
async def api_me(user: dict = Depends(get_current_user)):
    return {"code": 0, "data": user}


@app.get("/api/me/optional")
async def api_me_optional(user: Optional[dict] = Depends(get_optional_user)):
    if user:
        return {"code": 0, "data": user, "is_guest": False}
    return {"code": 0, "data": None, "is_guest": True}


@app.post("/api/logout")
async def api_logout():
    response = JSONResponse({"code": 0, "msg": "已登出"})
    response.delete_cookie(key="law_token", path="/", httponly=True, samesite="lax")
    return response


# ==================== 游客咨询接口 ====================

@app.post("/api/guest/consultation")
async def api_guest_consultation(request: Request, db: Session = Depends(get_db)):
    body = await request.json()
    question = body.get("question", "")
    guest_token = body.get("guest_token", "")
    if not question or not guest_token:
        raise HTTPException(400, "请输入咨询内容")
    gs = create_guest_session(guest_token, db)
    if gs.consult_count >= 10:
        return {"code": 1, "msg": "游客试用已满10轮，请注册账号后继续使用", "locked": True}
    kb_results = kb_search(question, n_results=3)
    kb_context = "\n".join([r["text"] for r in kb_results]) if kb_results else ""
    expert_ctx = get_expert_context(question)
    csr_prompt = "你是法务审核系统的客服专员。请根据知识库内容回答用户的法律咨询。回答要专业、简洁、有条理。"
    if expert_ctx:
        csr_prompt += f"\n\n可引用的外部法律专家：\n{expert_ctx}\n\n在回答中可根据需要引用上述专家意见。"
    user_msg = f"用户提问：{question}"
    if kb_context:
        user_msg += f"\n\n相关法条参考：\n{kb_context}"
    # Quick LLM availability check
    llm_cfg = get_llm_config()
    if not llm_cfg.get("api_key") and not llm_cfg.get("synced"):
        return {"code": 1, "msg": "LLM模型未配置，请先在管理后台检查模型同步状态"}
    try:
        model_name = body.get("model") or None
        answer = await asyncio.wait_for(call_llm(csr_prompt, user_msg, model_name), timeout=30)
    except asyncio.TimeoutError:
        return {"code": 1, "msg": "LLM响应超时，请稍后再试"}
    except Exception as e:
        import traceback; traceback.print_exc()
        return {"code": 1, "msg": f"LLM服务暂不可用: {str(e)[:200]}"}
    gs.consult_count += 1
    db.commit()
    return {"code": 0, "data": {"answer": answer, "remaining": 10 - gs.consult_count}, "locked": False}


@app.get("/api/guest/consult-count")
async def api_guest_consult_count(guest_token: str = Query(""), db: Session = Depends(get_db)):
    if not guest_token:
        return {"code": 0, "data": {"count": 0, "remaining": 10}}
    gs = db.query(GuestSession).filter(GuestSession.guest_token == guest_token).first()
    if not gs:
        return {"code": 0, "data": {"count": 0, "remaining": 10}}
    return {"code": 0, "data": {"count": gs.consult_count, "remaining": 10 - gs.consult_count}}


# ==================== 法务咨询接口（已登录） ====================

@app.post("/api/consultation")
async def api_consultation(request: Request, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    body = await request.json()
    question = body.get("question", "")
    if not question:
        raise HTTPException(400, "请输入咨询内容")
    kb_results = kb_search(question, n_results=3)
    kb_context = "\n".join([r["text"] for r in kb_results]) if kb_results else ""
    expert_ctx = get_expert_context(question)
    csr_prompt = "你是法务审核系统的客服专员。请根据知识库内容回答用户的法律咨询。回答要专业、简洁、有条理。"
    if expert_ctx:
        csr_prompt += f"\n\n可引用的外部法律专家：\n{expert_ctx}\n\n在回答中可根据需要引用上述专家意见。"
    user_msg = f"用户提问：{question}"
    if kb_context:
        user_msg += f"\n\n相关法条参考：\n{kb_context}"
    # Quick LLM availability check
    llm_cfg = get_llm_config()
    if not llm_cfg.get("api_key") and not llm_cfg.get("synced"):
        return {"code": 1, "msg": "LLM模型未配置，请先在管理后台检查模型同步状态"}
    try:
        model_name = body.get("model") or None
        answer = await asyncio.wait_for(call_llm(csr_prompt, user_msg, model_name), timeout=30)
    except asyncio.TimeoutError:
        return {"code": 1, "msg": "LLM响应超时，请稍后再试"}
    except Exception as e:
        import traceback; traceback.print_exc()
        return {"code": 1, "msg": f"LLM服务暂不可用: {str(e)[:200]}"}
    need_dispatch = any(kw in answer for kw in ["建议咨询", "需要专业", "已转派", "复杂"])
    work_order_id = None
    if need_dispatch:
        order = create_order(db, user["user_id"], "law", f"法律咨询: {question[:50]}", question, user["user_id"])
        work_order_id = order.id
        answer += f"\n\n[系统提示] 已生成工单#{order.id}，法务专员将尽快处理。"
    record = Consultation(user_id=user["user_id"], question=question, answer=answer, role_used="csr", work_order_id=work_order_id)
    db.add(record)
    db.commit()
    return {"code": 0, "data": {"answer": answer, "work_order_id": work_order_id}}
# ==================== 合同接口 ====================

@app.post("/api/contract/upload")
async def api_upload_contract(file: UploadFile = File(...), user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        return {"code": 1, "msg": "不支持的文件格式"}
    uid_dir = USER_CONTRACT_DIR / str(user["user_id"])
    uid_dir.mkdir(parents=True, exist_ok=True)
    file_path = uid_dir / file.filename
    with open(file_path, "wb") as f:
        f.write(await file.read())
    contract = Contract(user_id=user["user_id"], filename=file.filename, original_path=str(file_path), status="uploaded")
    db.add(contract)
    db.commit()
    db.refresh(contract)
    create_order(db, user["user_id"], "audit", f"合同审核: {file.filename}", f"合同ID: {contract.id}", user["user_id"])
    return {"code": 0, "msg": "上传成功，已创建审核工单", "data": {"contract_id": contract.id}}


@app.get("/api/contract/{contract_id}/status")
async def api_contract_status(contract_id: int, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    contract = db.query(Contract).filter(Contract.id == contract_id, Contract.user_id == user["user_id"]).first()
    if not contract:
        raise HTTPException(404, "合同不存在")
    return {"code": 0, "data": {"status": contract.status, "risk_summary": contract.risk_summary}}



@app.get("/api/contract/{contract_id}/risks")
async def api_get_contract_risks(contract_id: int, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """获取已审核合同的风险数据（不重新审核，直接读取存储的结果）"""
    contract = db.query(Contract).filter(Contract.id == contract_id, Contract.user_id == user["user_id"]).first()
    if not contract:
        raise HTTPException(404, "合同不存在")
    if contract.status != "audited":
        return {"code": 1, "msg": "合同尚未审核完成", "data": None}
    risks = []
    if contract.risk_summary:
        try:
            risks = json.loads(contract.risk_summary)
        except:
            risks = []
    return {"code": 0, "data": {
        "contract_id": contract.id,
        "filename": contract.filename,
        "status": contract.status,
        "risks": risks,
        "created_at": str(contract.created_at)
    }}

@app.get("/api/contract/{cid}/download-docx")
async def api_download_docx(cid: int, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """下载带批注的合同审核文档：风险点位置直接标注+Word批注+风险汇总表"""
    contract = db.query(Contract).filter(Contract.id == cid, Contract.user_id == user["user_id"]).first()
    if not contract:
        raise HTTPException(404, "合同不存在")
    risks = []
    if contract.risk_summary:
        try:
            risks = json.loads(contract.risk_summary)
        except:
            risks = []
    output_path = str(SAVE_DIR / f"annotated_contract_{cid}.docx")
    try:
        file_path = contract.original_path if (contract.original_path and os.path.exists(contract.original_path)) else None
        generate_annotated_docx(file_path, risks, output_path)
    except Exception as e:
        # Fallback: create a basic document
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading("合同审核结果", level=0)
        doc.add_paragraph(f"文件名: {contract.filename}")
        doc.add_paragraph(f"审核时间: {contract.created_at}")
        doc.add_paragraph(f"风险点数: {len(risks)}")
        if risks:
            for idx, risk in enumerate(risks, 1):
                lv = risk.get("level", "low")
                label = {"high": "高风险", "medium": "中风险", "low": "低风险"}.get(lv, lv)
                doc.add_paragraph(f"{idx}. [{label}] {risk.get('description', '-')}")
                if risk.get("suggestion"):
                    doc.add_paragraph(f"   修改建议: {risk['suggestion']}")
        doc.save(output_path)
    return FileResponse(
        output_path,
        filename=f"{contract.filename.replace('.','_')}_审核批注版.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.post("/api/contract/{contract_id}/audit")
async def api_audit_contract(contract_id: int, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(404, "合同不存在")
    contract.status = "auditing"
    db.commit()
    text = parse_contract(contract.original_path)
    analysis = await analyze_contract(text)
    report_path = await generate_report(contract_id, analysis)
    contract.status = "audited"
    contract.audit_report_path = report_path
    db.commit()
    record = AuditRecord(contract_id=contract_id, auditor_id=user["user_id"],
                         findings=analysis.get("risk_points", []),
                         risk_level=analysis.get("risk_level", "unknown"),
                         report_path=report_path)
    db.add(record)
    db.commit()
    risk_points = analysis.get("risk_points", [])
    normalized_risks = []
    for rp in risk_points:
        sev = rp.get("severity", rp.get("level", "low"))
        sev_map = {"critical": "high", "high": "high", "medium": "medium", "low": "low"}
        normalized_risks.append({
            "level": sev_map.get(sev, sev),
            "position": rp.get("position", rp.get("clause_text", "")[:30]),
            "clause": rp.get("clause_text", rp.get("clause", "")),
            "description": rp.get("risk_type", rp.get("description", "")),
            "law_basis": rp.get("law_basis", ""),
            "suggestion": rp.get("suggestion", "")
        })
    analysis["risks"] = normalized_risks
    contract.risk_summary = json.dumps(normalized_risks, ensure_ascii=False)
    db.commit()
    analysis["filename"] = contract.filename
    return {"code": 0, "msg": "审核完成", "data": analysis}


# ==================== 工单接口 ====================

@app.get("/api/workorders")
async def api_workorders(status: Optional[str] = None, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    role = user.get("role", "user")
    if role == "admin":
        orders = get_orders(db, status=status)
    else:
        orders = get_orders(db, role=role, status=status, user_id=user["user_id"])
    return {"code": 0, "data": [{"id": o.id, "title": o.title, "description": o.description, "status": o.status, "assigned_role": o.assigned_role, "result": o.result, "created_at": str(o.created_at)} for o in orders]}

@app.post("/api/workorder/create")
async def api_create_order(request: Request, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    body = await request.json()
    order = create_order(db, user["user_id"], body["role"], body["title"], body.get("description", ""), user["user_id"])
    return {"code": 0, "msg": "工单已创建", "data": {"id": order.id}}

@app.put("/api/workorder/{order_id}/accept")
async def api_accept_order(order_id: int, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    accept_order(db, order_id, user["user_id"])
    return {"code": 0, "msg": "已接单"}

@app.put("/api/workorder/{order_id}/complete")
async def api_complete_order(order_id: int, request: Request, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    body = await request.json()
    complete_order(db, order_id, body.get("result", ""))
    return {"code": 0, "msg": "已完成"}

@app.put("/api/workorder/{order_id}/return")
async def api_return_order(order_id: int, request: Request, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    body = await request.json()
    return_order(db, order_id, body.get("reason", ""))
    return {"code": 0, "msg": "已退回"}


# ==================== 历史记录接口 ====================

@app.get("/api/history")
async def api_history(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    consultations = db.query(Consultation).filter(Consultation.user_id == user["user_id"]).order_by(Consultation.created_at.desc()).limit(50).all()
    contracts = db.query(Contract).filter(Contract.user_id == user["user_id"]).order_by(Contract.created_at.desc()).limit(50).all()
    return {"code": 0, "data": {
        "consultations": [{"id": c.id, "question": c.question, "answer": c.answer, "created_at": str(c.created_at)} for c in consultations],
        "contracts": [{"id": c.id, "filename": c.filename, "status": c.status, "risk_summary": c.risk_summary, "created_at": str(c.created_at)} for c in contracts]
    }}


# ==================== 组织架构管理（admin专属） ====================

@app.get("/api/admin/org/tree")
async def api_org_tree(user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    nodes = db.query(Organization).filter(Organization.disabled == False).all()
    tree = []
    node_map = {}
    for n in nodes:
        node_map[n.id] = {"id": n.id, "type": n.node_type, "name": n.name, "parent_id": n.parent_id, "disabled": n.disabled, "children": []}
    for n in nodes:
        if n.parent_id and n.parent_id in node_map:
            node_map[n.parent_id]["children"].append(node_map[n.id])
        else:
            tree.append(node_map[n.id])
    return {"code": 0, "data": tree}


@app.post("/api/admin/org/add")
async def api_org_add(request: Request, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    body = await request.json()
    node_type = body.get("type", "")
    name = body.get("name", "").strip()
    parent_id = body.get("parent_id")
    if not node_type or not name:
        raise HTTPException(400, "类型和名称不能为空")
    node = Organization(node_type=node_type, name=name, parent_id=parent_id)
    db.add(node)
    db.commit()
    db.refresh(node)
    return {"code": 0, "msg": "添加成功", "data": {"id": node.id}}


@app.get("/api/admin/org/template")
async def api_org_template():
    """下载组织架构导入模板"""
    import io
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "组织架构"
    ws.append(["公司名称", "部门名称", "人员姓名"])
    ws.append(["示例科技公司", "技术部", "张三"])
    ws.append(["示例科技公司", "法务部", "李四"])
    ws.append(["示例科技公司", "法务部", "王五"])
    ws.append(["示例集团", "市场部", "赵六"])
    ws.column_dimensions["A"].width = 20
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 12
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    from fastapi.responses import StreamingResponse
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             headers={"Content-Disposition": "attachment; filename=org_template.xlsx"})


@app.post("/api/admin/org/upload")
async def api_org_upload(file: UploadFile = File(...), user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    """批量导入组织架构：Excel/CSV，格式：公司名称 | 部门名称 | 人员姓名"""
    import io, csv
    rows = []
    fname = (file.filename or "").lower()
    raw = await file.read()

    # 解析Excel文件
    if fname.endswith((".xlsx", ".xls")):
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(raw), read_only=True)
        ws = wb.active
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                continue  # 跳过表头
            vals = [str(c).strip() if c else "" for c in row]
            if len(vals) >= 3 and vals[0]:
                rows.append(vals[:3])
    # 解析CSV文件
    elif fname.endswith(".csv"):
        text = raw.decode("utf-8-sig")
        reader = csv.reader(io.StringIO(text))
        next(reader, None)  # 跳过表头
        for row in reader:
            vals = [c.strip() for c in row]
            if len(vals) >= 3 and vals[0]:
                rows.append(vals[:3])
    else:
        return {"code": -1, "msg": "仅支持 .xlsx / .csv 格式"}

    if not rows:
        return {"code": -1, "msg": "文件为空或格式不正确"}

    # 已有节点缓存：(type, name, parent_id) -> id
    existing = db.query(Organization).all()
    cache = {}
    for n in existing:
        cache[(n.node_type, n.name, n.parent_id)] = n.id

    added = {"company": 0, "department": 0, "person": 0}

    def get_or_create(node_type, name, parent_id):
        key = (node_type, name, parent_id)
        if key in cache:
            return cache[key], False
        node = Organization(node_type=node_type, name=name, parent_id=parent_id)
        db.add(node)
        db.flush()
        cache[key] = node.id
        return node.id, True

    for company, dept, person in rows:
        cid, c_new = get_or_create("company", company, None)
        if c_new:
            added["company"] += 1
        did, d_new = get_or_create("department", dept, cid)
        if d_new:
            added["department"] += 1
        if person:
            pid, p_new = get_or_create("person", person, did)
            if p_new:
                added["person"] += 1

    db.commit()
    total = sum(added.values())
    return {"code": 0, "msg": f"导入完成：新增公司{added['company']}个、部门{added['department']}个、人员{added['person']}个，共{total}条",
            "data": added}


@app.put("/api/admin/org/{node_id}/disable")
async def api_org_disable(node_id: int, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    node = db.query(Organization).filter(Organization.id == node_id).first()
    if not node:
        raise HTTPException(404, "节点不存在")
    node.disabled = not node.disabled
    db.commit()
    return {"code": 0, "msg": "已更新状态"}


@app.delete("/api/admin/org/{node_id}")
async def api_org_delete(node_id: int, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    node = db.query(Organization).filter(Organization.id == node_id).first()
    if not node:
        raise HTTPException(404, "节点不存在")
    db.delete(node)
    db.commit()
    return {"code": 0, "msg": "已删除"}


@app.get("/api/org/options")
async def api_org_options(db: Session = Depends(get_db)):
    nodes = db.query(Organization).filter(Organization.disabled == False).all()
    companies = [{"id": n.id, "name": n.name} for n in nodes if n.node_type == "company"]
    departments = [{"id": n.id, "name": n.name, "parent_id": n.parent_id} for n in nodes if n.node_type == "department"]
    persons = [{"id": n.id, "name": n.name, "parent_id": n.parent_id} for n in nodes if n.node_type == "person"]
    return {"code": 0, "data": {"companies": companies, "departments": departments, "persons": persons}}


# ==================== 账号管理（admin专属） ====================

@app.get("/api/admin/users")
async def api_admin_users(user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    users = db.query(User).all()
    return {"code": 0, "data": [{
        "id": u.id, "username": u.username, "role": u.role,
        "company": u.company or "", "department": u.department or "", "full_name": u.full_name or "",
        "enabled": u.enabled, "last_login": str(u.last_login)
    } for u in users]}


@app.post("/api/admin/user/create")
async def api_admin_create_user(request: Request, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    body = await request.json()
    result = register(body["username"], body["password"], db,
                      company=body.get("company", ""), department=body.get("department", ""),
                      full_name=body.get("full_name", ""))
    return {"code": 0, "msg": "账号创建成功", "data": result}


@app.put("/api/admin/user/{user_id}/reset-password")
async def api_admin_reset_password(user_id: int, request: Request, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    body = await request.json()
    target = db.query(User).filter(User.id == user_id).first()
    if not target:
        raise HTTPException(404, "用户不存在")
    target.password_hash = bcrypt.hashpw(body["password"].encode(), bcrypt.gensalt()).decode()
    db.commit()
    return {"code": 0, "msg": "密码已重置"}


@app.put("/api/admin/user/{user_id}/toggle")
async def api_admin_toggle_user(user_id: int, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    target = db.query(User).filter(User.id == user_id).first()
    if not target:
        raise HTTPException(404, "用户不存在")
    target.enabled = not target.enabled
    db.commit()
    return {"code": 0, "msg": "已更新", "enabled": target.enabled}


@app.put("/api/admin/user/{user_id}/update-dept")
async def api_admin_update_dept(user_id: int, request: Request, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    body = await request.json()
    target = db.query(User).filter(User.id == user_id).first()
    if not target:
        raise HTTPException(404, "用户不存在")
    target.company = body.get("company", target.company)
    target.department = body.get("department", target.department)
    db.commit()
    return {"code": 0, "msg": "部门已更新"}


# ==================== 知识库管理 ====================

@app.post("/api/admin/kb/upload")
async def api_admin_kb_upload(file: UploadFile = File(...), user: dict = Depends(require_admin)):
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        return {"code": 1, "msg": "不支持的文件格式"}
    file_path = KB_UPLOAD_DIR / file.filename
    with open(file_path, "wb") as f:
        f.write(await file.read())
    if ext == ".pdf":
        text = parse_pdf(str(file_path))
    elif ext == ".docx":
        text = parse_docx(str(file_path))
    elif ext == ".txt":
        text = file_path.read_text(encoding="utf-8")
    else:
        return {"code": 1, "msg": "不支持的文件格式"}
    add_document(text, {"source": file.filename, "type": "admin_upload"})
    return {"code": 0, "msg": "上传成功，已入库"}

@app.post("/api/admin/kb/rebuild")
async def api_admin_kb_rebuild(user: dict = Depends(require_admin)):
    await asyncio.to_thread(rebuild_index)
    return {"code": 0, "msg": "索引重建完成"}


# ==================== 模型状态 & 设备管理 ====================

@app.get("/api/dashboard")
async def api_dashboard(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """法务仪表盘数据聚合（多维度）"""
    from collections import Counter, defaultdict
    uid = user["user_id"]
    role = user.get("role", "user")
    # Admin sees all contracts, regular users see only their own
    if role == "admin":
        contracts = db.query(Contract).order_by(Contract.created_at.desc()).all()
        consultations = db.query(Consultation).order_by(Consultation.created_at.desc()).all()
    else:
        contracts = db.query(Contract).filter(Contract.user_id == uid).order_by(Contract.created_at.desc()).all()
        consultations = db.query(Consultation).filter(Consultation.user_id == uid).order_by(Consultation.created_at.desc()).all()

    total = len(contracts)
    audited = sum(1 for c in contracts if c.status == "audited")
    pending = total - audited
    high_risk = 0
    risk_types = Counter()
    monthly = Counter()
    recent_contracts = []
    recent_activities = []

    for c in contracts:
        # Risk analysis
        if c.risk_summary:
            try:
                risks = json.loads(c.risk_summary)
                if any(r.get("level") == "high" for r in risks):
                    high_risk += 1
                for r in risks:
                    rt = r.get("description", "").strip()
                    if rt:
                        risk_types[rt] += 1
            except:
                pass

        # Monthly stats
        if c.created_at:
            month_key = c.created_at.strftime("%Y-%m")
            monthly[month_key] += 1

        # Recent contracts
        if len(recent_contracts) < 8:
            recent_contracts.append({
                "id": c.id, "filename": c.filename,
                "status": c.status, "created_at": str(c.created_at)
            })

        # Activity feed
        if len(recent_activities) < 15:
            recent_activities.append({
                "type": "upload" if c.status == "uploaded" else "audit_done" if c.status == "audited" else "audit",
                "desc": f"上传合同 {c.filename}" if c.status == "uploaded" else f"完成审核 {c.filename}",
                "time": str(c.created_at)[:16]
            })

    for co in consultations[:10]:
        if len(recent_activities) < 20:
            recent_activities.append({
                "type": "consult",
                "desc": f"法务咨询: {co.question[:30]}...",
                "time": str(co.created_at)[:16]
            })

    recent_activities.sort(key=lambda x: x["time"], reverse=True)
    recent_activities = recent_activities[:15]

    # Contract type guess based on filename
    type_map = {"租赁": 0, "采购": 0, "劳务": 0, "服务": 0, "销售": 0, "其他": 0}
    for c in contracts:
        name = c.filename
        matched = False
        for t in type_map:
            if t in name:
                type_map[t] += 1
                matched = True
                break
        if not matched:
            type_map["其他"] += 1

    return {"code": 0, "data": {
        "total_contracts": total,
        "audited_contracts": audited,
        "high_risk_contracts": high_risk,
        "pending_contracts": pending,
        "total_consultations": len(consultations),
        "risk_types": [{"name": k, "count": v} for k, v in risk_types.most_common(8)],
        "monthly": [{"month": k, "count": v} for k, v in sorted(monthly.items())],
        "contract_types": [{"name": k, "count": v} for k, v in type_map.items() if v > 0],
        "recent_contracts": recent_contracts[:8],
        "recent_activities": recent_activities[:15]
    }}

@app.get("/api/admin/model-status")
async def api_model_status(user: dict = Depends(require_admin)):
    cfg = get_llm_config()
    return {"code": 0, "data": {"model": cfg.get("model", "未配置"), "synced": cfg.get("synced", False), "last_sync": cfg.get("last_sync"), "source": cfg.get("source"), "error": cfg.get("error")}}

@app.post("/api/admin/change-password")
async def api_change_password(request: Request, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    """管理员修改任意用户密码"""
    body = await request.json()
    user_id = body.get("user_id")
    new_password = body.get("new_password", "")
    if not user_id or not new_password or len(new_password) < 6:
        return {"code": 1, "msg": "请提供有效用户ID和密码（至少6位）"}
    target = db.query(User).filter(User.id == user_id).first()
    if not target:
        return {"code": 1, "msg": "用户不存在"}
    import bcrypt
    target.password_hash = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode()
    db.commit()
    return {"code": 0, "msg": f"用户 {target.username} 密码已更新"}

@app.post("/api/admin/switch-model")
async def api_switch_model(request: Request, user: dict = Depends(require_admin)):
    """切换全局LLM模型"""
    body = await request.json()
    model = body.get("model", "")
    if not model:
        return {"code": 1, "msg": "请指定模型名称"}
    from .llm_sync.sync import _cached_config
    _cached_config["model"] = model
    from .llm_sync.sync import set_manual_model
    set_manual_model(model)
    return {"code": 0, "msg": f"模型已切换至 {model}", "data": {"model": model}}

@app.post("/api/admin/device/approve/{binding_id}")
async def api_approve_device(binding_id: int, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    binding = db.query(DeviceBinding).filter(DeviceBinding.id == binding_id).first()
    if not binding:
        raise HTTPException(404, "记录不存在")
    binding.approved = True
    db.commit()
    return {"code": 0, "msg": "设备已批准"}

@app.get("/api/admin/devices")
async def api_admin_devices(user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    bindings = db.query(DeviceBinding).all()
    return {"code": 0, "data": [{"id": b.id, "user_id": b.user_id, "approved": b.approved, "created_at": str(b.created_at)} for b in bindings]}

@app.get("/api/admin/stats")
async def api_admin_stats(user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    stats = {}
    for role in ["csr", "law", "audit", "sort", "kb"]:
        total = db.query(WorkOrder).filter(WorkOrder.assigned_role == role).count()
        done = db.query(WorkOrder).filter(WorkOrder.assigned_role == role, WorkOrder.status == "done").count()
        stats[role] = {"total": total, "done": done}
    return {"code": 0, "data": stats}

@app.get("/api/admin/workorders")
async def api_admin_workorders(user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    orders = get_orders(db)
    return {"code": 0, "data": [{"id": o.id, "title": o.title, "status": o.status, "assigned_role": o.assigned_role, "user_id": o.user_id, "created_at": str(o.created_at)} for o in orders]}


# ==================== 启动入口 ====================

def main():
    import uvicorn
    uvicorn.run(app, host=HOST, port=PORT, log_level="info")

if __name__ == "__main__":
    main()
