"""
Word文档批注(评论)工具
在python-docx基础上增加对Word批注(Comments)的原生XML操作支持
"""
import os
from datetime import datetime
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor


def _get_or_create_comments_part(document):
    """获取或创建文档的comments部件"""
    doc_part = document.part
    comments_part = None
    comments_elem = None

    for rel in doc_part.rels.values():
        if rel.reltype.endswith("/comments"):
            comments_part = rel.target_part
            comments_elem = comments_part._element
            break

    if comments_elem is None:
        comments_xml = (
            '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" />'
        )
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
        partname = PackURI("/word/comments.xml")

        comments_elem = parse_xml(comments_xml)
        comments_part = Part(
            partname, content_type,
            etree.tostring(comments_elem, xml_declaration=True, encoding="UTF-8", standalone=True),
            doc_part.package
        )
        comments_part._element = comments_elem

        from docx.opc.rel import Rel
        rId = f"rId{len(doc_part.rels) + 100}"
        rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
        rel = Rel(doc_part, rId, rel_type, target=comments_part)
        doc_part.rels._rels[rId] = rel

    return comments_part, comments_elem


def add_comment_to_paragraph(paragraph, comment_text, author="法务审核系统"):
    """为指定段落添加Word批注（评论气泡），返回评论ID"""
    document = paragraph.part.document
    _, comments_elem = _get_or_create_comments_part(document)

    existing = comments_elem.findall(qn("w:comment"))
    comment_id = len(existing)

    comment = etree.SubElement(comments_elem, qn("w:comment"))
    comment.set(qn("w:id"), str(comment_id))
    comment.set(qn("w:author"), author)
    comment.set(qn("w:date"), datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"))

    cp = etree.SubElement(comment, qn("w:p"))
    pPr = etree.SubElement(cp, qn("w:pPr"))
    pStyle = etree.SubElement(pPr, qn("w:pStyle"))
    pStyle.set(qn("w:val"), "CommentText")

    lines = comment_text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            br = etree.SubElement(cp, qn("w:r"))
            etree.SubElement(br, qn("w:br"))
        if line.strip():
            r = etree.SubElement(cp, qn("w:r"))
            rPr = etree.SubElement(r, qn("w:rPr"))
            rStyle = etree.SubElement(rPr, qn("w:rStyle"))
            rStyle.set(qn("w:val"), "CommentReference")
            sz = etree.SubElement(rPr, qn("w:sz"))
            sz.set(qn("w:val"), "18")
            t = etree.SubElement(r, qn("w:t"))
            t.text = line
            t.set(qn("xml:space"), "preserve")

    p_elem = paragraph._element

    cr_start = etree.Element(qn("w:commentRangeStart"))
    cr_start.set(qn("w:id"), str(comment_id))
    p_elem.insert(0, cr_start)

    cr_end = etree.Element(qn("w:commentRangeEnd"))
    cr_end.set(qn("w:id"), str(comment_id))
    p_elem.addnext(cr_end)

    runs = p_elem.findall(qn("w:r"))
    if runs:
        last_run = runs[-1]
        comm_ref = etree.SubElement(last_run, qn("w:commentReference"))
        comm_ref.set(qn("w:id"), str(comment_id))
    else:
        r = etree.SubElement(p_elem, qn("w:r"))
        comm_ref = etree.SubElement(r, qn("w:commentReference"))
        comm_ref.set(qn("w:id"), str(comment_id))

    return comment_id


def _safe_heading(doc, text, level=1):
    """添加标题，若样式不存在则用加粗段落替代"""
    try:
        return doc.add_heading(text, level=level)
    except KeyError:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
        return p



def generate_annotated_docx(original_path, risks, output_path):
    """
    生成带全文标注的合同审核文档：
    解析合同全文并重新构建文档，在风险点位置插入风险描述和修改建议，
    保留完整原文，不依赖原docx样式
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    import re

    # Parse contract text from original file
    contract_text = ""
    if original_path and os.path.exists(original_path):
        try:
            from .law_kb.knowledge_base import parse_pdf, parse_docx
            ext = original_path.lower().split(".")[-1] if "." in original_path else ""
            if ext == "pdf":
                contract_text = parse_pdf(original_path)
            elif ext == "docx":
                contract_text = parse_docx(original_path)
            else:
                with open(original_path, "r", encoding="utf-8") as f:
                    contract_text = f.read()
        except:
            contract_text = ""
    if not contract_text or len(contract_text) < 20:
        contract_text = "(无法读取合同原文)"

    doc = Document()
    
    def add_h(text, level=1):
        try:
            return doc.add_heading(text, level=level)
        except KeyError:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(16 if level == 1 else 14)
            return p

    level_icons = {"high": "\U0001f534", "medium": "\U0001f7e1", "low": "\U0001f7e2"}
    level_labels = {"high": "\u9ad8\u98ce\u9669", "medium": "\u4e2d\u98ce\u9669", "low": "\u4f4e\u98ce\u9669"}
    level_colors = {
        "high": RGBColor(0xC6, 0x28, 0x28),
        "medium": RGBColor(0xEF, 0x6C, 0x00),
        "low": RGBColor(0x2E, 0x7D, 0x32),
    }

    # Section 1: Full contract text with inline annotations
    add_h("\u5408\u540c\u539f\u6587\uff08\u98ce\u9669\u6807\u6ce8\u7248\uff09", 1)
    doc.add_paragraph("")

    paragraphs = contract_text.split("\n")
    matched_set = set()

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # Write original paragraph
        p = doc.add_paragraph()
        run = p.add_run(para_text)
        run.font.size = Pt(11)

        # Try to match against risks
        matched = None
        for ri, risk in enumerate(risks):
            if ri in matched_set:
                continue
            clause = risk.get("clause", "") or ""
            position = risk.get("position", "") or ""
            # Multiple matching strategies
            match = False
            if clause:
                for sl in [25, 20, 15, 12, 10, 8]:
                    if len(clause) >= sl and clause[:sl] in para_text:
                        match = True
                        break
            if not match and position:
                for sl in [15, 12, 10, 8]:
                    if len(position) >= sl and position[:sl] in para_text:
                        match = True
                        break
            if match:
                matched = risk
                matched_set.add(ri)
                break

        if matched:
            lv = matched.get("level", "low")
            icon = level_icons.get(lv, "\u26aa")
            label = level_labels.get(lv, lv)
            desc = matched.get("description", "") or ""
            sugg = matched.get("suggestion", "") or ""
            law = matched.get("law_basis", "") or ""
            color = level_colors.get(lv, RGBColor(0, 0, 0))

            # Visible annotation block
            note = doc.add_paragraph()
            note.paragraph_format.space_before = Pt(2)
            note.paragraph_format.space_after = Pt(1)
            r = note.add_run(f"{icon} [{label}] {desc}")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = color

            if law:
                nl = doc.add_paragraph()
                nl.paragraph_format.space_before = Pt(0)
                nl.paragraph_format.space_after = Pt(1)
                rl = nl.add_run(f"\u2696\ufe0f \u6cd5\u5f8b\u4f9d\u636e\uff1a{law}")
                rl.font.size = Pt(10)
                rl.italic = True

            if sugg:
                ns = doc.add_paragraph()
                ns.paragraph_format.space_before = Pt(0)
                ns.paragraph_format.space_after = Pt(4)
                rs = ns.add_run(f"\U0001f4a1 \u4fee\u6539\u5efa\u8bae\uff1a{sugg}")
                rs.font.size = Pt(10)
                rs.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            
            # Separator
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(0)
            sep.paragraph_format.space_after = Pt(2)
            rs2 = sep.add_run("\u2500" * 40)
            rs2.font.size = Pt(8)
            rs2.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    # Section 2: Risk summary table
    doc.add_paragraph("")
    add_h("\U0001f4cb \u98ce\u9669\u5ba1\u6838\u6c47\u603b", 1)
    doc.add_paragraph(f"\u53d1\u73b0\u98ce\u9669\u70b9: {len(risks)} \u5904")
    
    high_c = sum(1 for r in risks if r.get("level") == "high")
    med_c = sum(1 for r in risks if r.get("level") == "medium")
    low_c = sum(1 for r in risks if r.get("level") == "low")
    doc.add_paragraph(f"\U0001f534 \u9ad8\u98ce\u9669: {high_c}   \U0001f7e1 \u4e2d\u98ce\u9669: {med_c}   \U0001f7e2 \u4f4e\u98ce\u9669: {low_c}")

    if risks:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for i, h in enumerate(["\u5e8f\u53f7", "\u7b49\u7ea7", "\u6761\u6b3e\u4f4d\u7f6e", "\u98ce\u9669\u63cf\u8ff0", "\u4fee\u6539\u5efa\u8bae"]):
            cell = table.rows[0].cells[i]
            cell.text = h
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        for idx, risk in enumerate(risks, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = level_labels.get(risk.get("level", "low"), risk.get("level", "low"))
            row.cells[2].text = (risk.get("position", "") or risk.get("clause", "")[:30] or "-")
            row.cells[3].text = risk.get("description", "-")
            sug = risk.get("suggestion", "")
            law = risk.get("law_basis", "")
            sug_text = ""
            if law:
                sug_text = f"[\u4f9d\u636e] {law}\n"
            if sug:
                sug_text += f"[\u5efa\u8bae] {sug}"
            row.cells[4].text = sug_text or "-"

    doc.save(output_path)
    return output_path

