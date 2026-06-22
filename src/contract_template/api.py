"""
Contract Template Management & Generation API
Smart field detection + precise fill-back (preserves original formatting)
"""
import os
import json
import shutil
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Request
from fastapi.responses import FileResponse, JSONResponse
from sqlalchemy.orm import Session

from ..config import TEMPLATE_DIR, GENERATED_CONTRACT_DIR, BASE_DIR
from ..database import get_db, ContractTemplate, GeneratedContract, User
from ..user_auth.auth import get_current_user, require_admin
from ..llm_sync.sync import call_llm
from .field_detector import detect_fields, replace_detected_fields

router = APIRouter(prefix="/api/contract-template", tags=["合同模板"])


# ==================== Template category list ====================

TEMPLATE_CATEGORIES = [
    "劳动合同", "采购合同", "销售合同", "租赁合同", "服务合同",
    "保密协议", "合作协议", "借款合同", "担保合同", "委托合同", "其他"
]


# ==================== Admin endpoints ====================

@router.get("/admin/categories")
async def get_categories(user: dict = Depends(require_admin)):
    return {"code": 0, "data": TEMPLATE_CATEGORIES}


@router.post("/admin/upload")
async def upload_template(
    name: str = Form(...),
    category: str = Form(...),
    description: str = Form(""),
    fill_fields: str = Form("[]"),
    auto_detect: str = Form("false"),
    file: UploadFile = File(...),
    user: dict = Depends(require_admin),
    db: Session = Depends(get_db)
):
    """Upload a contract template file (with optional smart field detection)."""
    # Save file
    cat_dir = TEMPLATE_DIR / category
    cat_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    safe_name = f"{ts}_{file.filename}"
    file_path = cat_dir / safe_name
    with open(file_path, "wb") as f:
        f.write(await file.read())

    # Extract template text
    template_text = _extract_template_text(str(file_path))

    # Parse fill fields (manual or auto-detected)
    try:
        fields = json.loads(fill_fields)
    except Exception:
        fields = []

    # Smart auto-detection: if enabled and no manual fields provided
    if auto_detect.lower() == "true" and not fields:
        doc = _load_docx(str(file_path))
        if doc:
            fields = detect_fields(doc)

    tpl = ContractTemplate(
        name=name, category=category, description=description,
        template_file_path=str(file_path), template_text=template_text,
        fill_fields=fields, created_by=user["user_id"]
    )
    db.add(tpl)
    db.commit()
    db.refresh(tpl)
    return {"code": 0, "msg": "模板上传成功", "data": {"id": tpl.id, "name": tpl.name}}


@router.post("/admin/{template_id}/auto-detect")
async def admin_auto_detect(template_id: int, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    """Admin: trigger smart field detection for a template."""
    tpl = db.query(ContractTemplate).filter(ContractTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(404, "模板不存在")
    doc = _load_docx(tpl.template_file_path)
    fields = []
    if doc:
        fields = detect_fields(doc)
    if not fields:
        fields = _get_fallback_fields(tpl.category or "其他")
    tpl.fill_fields = fields
    tpl.updated_at = datetime.utcnow()
    db.commit()
    return {"code": 0, "msg": f"识别完成，发现 {len(fields)} 个待填写字段", "data": fields}


@router.put("/admin/{template_id}/fill-fields")
async def update_fill_fields(template_id: int, request: Request, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    """Admin: manually update the fill fields for a template."""
    tpl = db.query(ContractTemplate).filter(ContractTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(404, "模板不存在")
    body = await request.json()
    fields = body.get("fields", [])
    tpl.fill_fields = fields
    tpl.updated_at = datetime.utcnow()
    db.commit()
    return {"code": 0, "msg": "字段已更新"}


@router.get("/admin/list")
async def list_templates(
    category: Optional[str] = None,
    user: dict = Depends(require_admin),
    db: Session = Depends(get_db)
):
    q = db.query(ContractTemplate)
    if category:
        q = q.filter(ContractTemplate.category == category)
    tpls = q.order_by(ContractTemplate.created_at.desc()).all()
    return {"code": 0, "data": [_tpl_to_dict(t) for t in tpls]}


@router.get("/admin/{template_id}")
async def get_template_detail(template_id: int, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    tpl = db.query(ContractTemplate).filter(ContractTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(404, "模板不存在")
    return {"code": 0, "data": _tpl_to_dict(tpl, detail=True)}


@router.put("/admin/{template_id}")
async def update_template(
    template_id: int,
    name: str = Form(""),
    category: str = Form(""),
    description: str = Form(""),
    fill_fields: str = Form(""),
    file: Optional[UploadFile] = File(None),
    user: dict = Depends(require_admin),
    db: Session = Depends(get_db)
):
    tpl = db.query(ContractTemplate).filter(ContractTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(404, "模板不存在")
    if name:
        tpl.name = name
    if category:
        tpl.category = category
    if description:
        tpl.description = description
    if fill_fields:
        try:
            tpl.fill_fields = json.loads(fill_fields)
        except Exception:
            pass
    if file:
        cat_dir = TEMPLATE_DIR / (category or tpl.category)
        cat_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d%H%M%S")
        file_path = cat_dir / f"{ts}_{file.filename}"
        with open(file_path, "wb") as f:
            f.write(await file.read())
        tpl.template_file_path = str(file_path)
        tpl.template_text = _extract_template_text(str(file_path))
    tpl.version += 1
    tpl.updated_at = datetime.utcnow()
    db.commit()
    return {"code": 0, "msg": "模板已更新"}


@router.put("/admin/{template_id}/toggle")
async def toggle_template(template_id: int, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    tpl = db.query(ContractTemplate).filter(ContractTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(404, "模板不存在")
    tpl.status = "disabled" if tpl.status == "active" else "active"
    db.commit()
    return {"code": 0, "msg": f"模板已{'停用' if tpl.status == 'disabled' else '启用'}"}


@router.delete("/admin/{template_id}")
async def delete_template(template_id: int, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    tpl = db.query(ContractTemplate).filter(ContractTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(404, "模板不存在")
    db.delete(tpl)
    db.commit()
    return {"code": 0, "msg": "模板已删除"}


@router.get("/admin/{template_id}/download")
async def download_template_file(template_id: int, user: dict = Depends(require_admin), db: Session = Depends(get_db)):
    tpl = db.query(ContractTemplate).filter(ContractTemplate.id == template_id).first()
    if not tpl or not tpl.template_file_path or not os.path.exists(tpl.template_file_path):
        raise HTTPException(404, "模板文件不存在")
    return FileResponse(tpl.template_file_path, filename=os.path.basename(tpl.template_file_path))


# ==================== User endpoints ====================

@router.get("/categories")
async def user_categories(user: dict = Depends(get_current_user)):
    return {"code": 0, "data": TEMPLATE_CATEGORIES}


@router.get("/list")
async def user_list_templates(
    category: Optional[str] = None,
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    q = db.query(ContractTemplate).filter(ContractTemplate.status == "active")
    if category:
        q = q.filter(ContractTemplate.category == category)
    tpls = q.order_by(ContractTemplate.category, ContractTemplate.name).all()
    return {"code": 0, "data": [_tpl_to_dict(t) for t in tpls]}


@router.get("/{template_id}/fields")
async def get_template_fields(template_id: int, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get template fill fields (pre-defined or auto-detected)."""
    tpl = db.query(ContractTemplate).filter(
        ContractTemplate.id == template_id, ContractTemplate.status == "active"
    ).first()
    if not tpl:
        raise HTTPException(404, "模板不存在")
    return {"code": 0, "data": {
        "id": tpl.id, "name": tpl.name, "category": tpl.category,
        "description": tpl.description, "fields": tpl.fill_fields or []
    }}


@router.get("/{template_id}/auto-fields")
async def get_auto_fields(template_id: int, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    """
    Get auto-detected fill fields for a template.
    Returns stored fields if available; otherwise runs smart detection on-the-fly.
    """
    tpl = db.query(ContractTemplate).filter(
        ContractTemplate.id == template_id, ContractTemplate.status == "active"
    ).first()
    if not tpl:
        raise HTTPException(404, "模板不存在")

    # Return stored fields if already detected
    if tpl.fill_fields:
        return {"code": 0, "data": {
            "id": tpl.id, "name": tpl.name, "category": tpl.category,
            "description": tpl.description, "fields": tpl.fill_fields
        }}

    # Run smart detection on-the-fly
    doc = _load_docx(tpl.template_file_path)
    fields = []
    if doc:
        fields = detect_fields(doc)

    # Fallback: if no fields detected, generate standard contract fields
    if not fields:
        fields = _get_fallback_fields(tpl.category or "其他")

    # Store for future use
    tpl.fill_fields = fields
    db.commit()

    return {"code": 0, "data": {
        "id": tpl.id, "name": tpl.name, "category": tpl.category,
        "description": tpl.description, "fields": fields
    }}


@router.post("/generate")
async def generate_contract(
    request: Request,
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Generate a contract: smart fill-back into the original template (preserves formatting)."""
    body = await request.json()
    template_id = body.get("template_id")
    filled_data = body.get("filled_data", {})
    title = body.get("title", "")

    tpl = db.query(ContractTemplate).filter(
        ContractTemplate.id == template_id, ContractTemplate.status == "active"
    ).first()
    if not tpl:
        raise HTTPException(404, "模板不存在或已停用")

    output_title = title or f"{tpl.name}_{datetime.now().strftime('%Y%m%d%H%M%S')}"

    # Strategy 1: Smart fill-back (preserve original template formatting)
    output_path = _smart_fill_back(tpl, filled_data, output_title, db)

    # Strategy 2 (fallback): LLM generation (if fill-back not possible)
    if not output_path:
        contract_text = await _generate_contract_text(tpl, filled_data)
        output_path = await _save_as_docx(contract_text, output_title, tpl.name)

    record = GeneratedContract(
        template_id=tpl.id, user_id=user["user_id"],
        title=output_title, filled_data=filled_data,
        output_path=str(output_path), status="generated"
    )
    db.add(record)
    db.commit()
    db.refresh(record)

    return {"code": 0, "msg": "合同生成成功", "data": {
        "id": record.id, "title": output_title,
        "download_url": f"/api/contract-template/download/{record.id}"
    }}


@router.get("/download/{record_id}")
async def download_generated(record_id: int, user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    record = db.query(GeneratedContract).filter(
        GeneratedContract.id == record_id,
        GeneratedContract.user_id == user["user_id"]
    ).first()
    if not record or not record.output_path or not os.path.exists(record.output_path):
        raise HTTPException(404, "合同文件不存在")
    record.status = "downloaded"
    db.commit()
    return FileResponse(
        record.output_path,
        filename=f"{record.title}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.get("/history")
async def user_history(user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    records = db.query(GeneratedContract).filter(
        GeneratedContract.user_id == user["user_id"]
    ).order_by(GeneratedContract.created_at.desc()).limit(50).all()
    return {"code": 0, "data": [{
        "id": r.id, "template_id": r.template_id, "title": r.title,
        "status": r.status, "created_at": str(r.created_at)
    } for r in records]}


# ==================== Internal helper functions ====================

def _convert_doc_to_docx(doc_path: str) -> typing.Optional[str]:
    """Convert .doc to .docx. Tries: LibreOffice headless > doc2docx > win32com."""
    import subprocess
    parent = str(Path(doc_path).parent)
    # Method 1: LibreOffice headless (cross-platform)
    for lo_cmd in ["libreoffice", "soffice"]:
        try:
            r = subprocess.run([lo_cmd, "--headless", "--convert-to", "docx", doc_path, "--outdir", parent],
                               capture_output=True, text=True, timeout=60)
            if r.returncode == 0:
                expected = str(Path(doc_path).with_suffix(".docx"))
                if os.path.exists(expected):
                    return expected
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    # Method 2: doc2docx package
    try:
        from doc2docx import convert
        out = convert(doc_path)
        if out and os.path.exists(out):
            return out
    except (ImportError, Exception):
        pass
    # Method 3: win32com (Windows only)
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(Path(doc_path).resolve()))
        out_path = str(Path(doc_path).with_suffix(".docx"))
        doc.SaveAs2(out_path, FileFormat=16)
        doc.Close()
        word.Quit()
        if os.path.exists(out_path):
            return out_path
    except Exception:
        pass
    return None


def _load_docx(file_path: str):
    """Load a .docx file as a python-docx Document.
    For .doc files, attempts automatic conversion to .docx first.
    Returns None if not supported or conversion fails."""
    if not file_path or not os.path.exists(file_path):
        return None
    ext = Path(file_path).suffix.lower()
    if ext == ".docx":
        try:
            from docx import Document
            return Document(file_path)
        except Exception:
            return None
    if ext == ".doc":
        converted = _convert_doc_to_docx(file_path)
        if converted:
            try:
                from docx import Document
                return Document(converted)
            except Exception:
                return None
        return None
    return None


def _extract_template_text(file_path: str) -> str:
    """Extract plain text from template file (including table text)."""
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            # Include both paragraph and table text for richer extraction
            lines = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            lines.append(cell_text)
            return "\n".join(lines)
        elif ext == ".pdf":
            import fitz
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            return text
        elif ext == ".txt":
            return Path(file_path).read_text(encoding="utf-8")
        elif ext == ".doc":
            # .doc format: try conversion then extract
            import subprocess
            converted = _convert_doc_to_docx(file_path)
            if converted:
                try:
                    from docx import Document as _Doc
                    d = _Doc(converted)
                    lines = [p.text for p in d.paragraphs]
                    for table in d.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    lines.append(cell.text.strip())
                    return "\n".join(lines)
                except Exception:
                    pass
            for cmd in ['antiword', 'catdoc']:
                try:
                    r = subprocess.run([cmd, file_path], capture_output=True, text=True, timeout=10)
                    if r.returncode == 0 and r.stdout.strip():
                        return r.stdout
                except Exception:
                    continue
            return "[.doc 格式：已存储，自动识别可能受限，建议转为 .docx 获取最佳效果]"
    except Exception as e:
        return f"[解析失败: {e}]"
    return ""


def _smart_fill_back(tpl, filled_data: dict, title: str, db) -> str:
    """
    Smart fill-back: load the original template .docx, replace detected field
    positions with user values, save as a new .docx.
    Preserves ALL original formatting, styles, tables, and structure.
    Returns the output file path, or None if fill-back is not possible.
    """
    if not tpl.template_file_path or not os.path.exists(tpl.template_file_path):
        return None
    ext = Path(tpl.template_file_path).suffix.lower()
    template_path = tpl.template_file_path
    if ext == ".doc":
        converted = _convert_doc_to_docx(tpl.template_file_path)
        if not converted:
            return None
        template_path = converted
    elif ext != ".docx":
        return None

    # Get fields (from DB or auto-detect)
    fields = tpl.fill_fields
    if not fields:
        doc = _load_docx(tpl.template_file_path)
        if not doc:
            return None
        fields = detect_fields(doc)
        if fields:
            tpl.fill_fields = fields
            db.commit()
    if not fields:
        return None

    try:
        from docx import Document
        doc = Document(template_path)

        # Replace detected field positions with user values
        replace_detected_fields(doc, fields, filled_data)

        # Save
        ts = datetime.now().strftime("%Y%m%d%H%M%S")
        safe_title = "".join(c for c in title if c.isalnum() or c in "._- ")
        filename = f"{safe_title}_{ts}.docx"
        output_path = str(GENERATED_CONTRACT_DIR / filename)
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"[Smart fill-back failed] {e}")
        return None


async def _generate_contract_text(tpl: ContractTemplate, filled_data: dict) -> str:
    """Generate complete contract text via LLM (fallback when fill-back is not possible)."""
    fields_desc = ""
    if tpl.fill_fields:
        for f in tpl.fill_fields:
            key = f.get("key", "")
            label = f.get("label", key)
            val = filled_data.get(key, "")
            fields_desc += f"  - {label}（{key}）：{val}\n"
    else:
        for k, v in filled_data.items():
            fields_desc += f"  - {k}：{v}\n"

    system_prompt = """你是一位专业法务文书撰写专家。根据提供的合同模板和用户填写的业务信息，生成一份完整、规范、合规的合同文本。

要求：
1. 严格遵循模板的结构和条款框架
2. 将用户提供的信息准确填入对应位置
3. 对于用户未提供但合同必需的内容，用【待填写】标注
4. 保持法律用语的专业性和规范性
5. 输出纯文本格式，保留条款编号和层级结构
6. 不要输出任何解释说明，直接输出合同正文"""

    user_msg = f"""【合同模板信息】
模板名称：{tpl.name}
模板分类：{tpl.category}
模板说明：{tpl.description}

【模板原文】
{tpl.template_text[:6000]}

【用户填写的业务信息】
{fields_desc}

请根据以上模板和业务信息，生成完整合同文本。"""

    result = await call_llm(system_prompt, user_msg)
    return result


async def _save_as_docx(contract_text: str, title: str, template_name: str) -> str:
    """Save LLM-generated contract text as a .docx file."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"（基于模板：{template_name}）")
    run.font.size = Pt(10)

    doc.add_paragraph("")

    for line in contract_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("第") and ("条" in line[:8] or "章" in line[:8]):
            doc.add_heading(line, level=2)
        elif line.startswith("甲方") or line.startswith("乙方") or line.startswith("丙方"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            doc.add_paragraph(line)

    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    safe_title = "".join(c for c in title if c.isalnum() or c in "._- ")
    filename = f"{safe_title}_{ts}.docx"
    output_path = str(GENERATED_CONTRACT_DIR / filename)
    doc.save(output_path)
    return output_path


def _get_fallback_fields(category: str) -> list:
    """Generate standard contract fields when auto-detection finds nothing.
    Ensures the form is never blank -- provides category-specific defaults."""
    base_fields = [
        {"key": "contract_no", "label": "合同编号", "type": "text", "group": "基本信息", "priority": 1},
        {"key": "party_a_name", "label": "甲方（名称）", "type": "text", "group": "甲乙双方信息", "priority": 1},
        {"key": "party_a_rep", "label": "甲方法定代表人", "type": "text", "group": "甲乙双方信息", "priority": 1},
        {"key": "party_a_addr", "label": "甲方地址", "type": "text", "group": "甲乙双方信息", "priority": 2},
        {"key": "party_a_phone", "label": "甲方联系电话", "type": "text", "group": "甲乙双方信息", "priority": 2},
        {"key": "party_b_name", "label": "乙方（名称）", "type": "text", "group": "甲乙双方信息", "priority": 1},
        {"key": "party_b_rep", "label": "乙方法定代表人", "type": "text", "group": "甲乙双方信息", "priority": 1},
        {"key": "party_b_addr", "label": "乙方地址", "type": "text", "group": "甲乙双方信息", "priority": 2},
        {"key": "party_b_phone", "label": "乙方联系电话", "type": "text", "group": "甲乙双方信息", "priority": 2},
        {"key": "total_amount", "label": "合同总金额（元）", "type": "number", "group": "金额信息", "priority": 4},
        {"key": "sign_date", "label": "签订日期", "type": "date", "group": "日期信息", "priority": 3},
        {"key": "effective_date", "label": "生效日期", "type": "date", "group": "日期信息", "priority": 3},
        {"key": "expire_date", "label": "截止日期", "type": "date", "group": "日期信息", "priority": 3},
        {"key": "sign_location", "label": "签订地点", "type": "text", "group": "基本信息", "priority": 5},
    ]
    category_extras = {
        "劳动合同": [
            {"key": "employee_name", "label": "员工姓名", "type": "text", "group": "甲乙双方信息", "priority": 1},
            {"key": "position", "label": "岗位/职位", "type": "text", "group": "合同内容", "priority": 5},
            {"key": "salary", "label": "薪资待遇", "type": "number", "group": "金额信息", "priority": 4},
            {"key": "work_location", "label": "工作地点", "type": "text", "group": "合同内容", "priority": 5},
        ],
        "采购合同": [
            {"key": "goods_name", "label": "采购物品名称", "type": "text", "group": "合同内容", "priority": 5},
            {"key": "quantity", "label": "数量", "type": "text", "group": "合同内容", "priority": 5},
            {"key": "unit_price", "label": "单价", "type": "number", "group": "金额信息", "priority": 4},
            {"key": "delivery_date", "label": "交货日期", "type": "date", "group": "日期信息", "priority": 3},
            {"key": "delivery_location", "label": "交货地点", "type": "text", "group": "合同内容", "priority": 5},
        ],
        "服务合同": [
            {"key": "service_content", "label": "服务内容", "type": "textarea", "group": "合同内容", "priority": 6},
            {"key": "service_period", "label": "服务期限", "type": "text", "group": "合同内容", "priority": 5},
            {"key": "service_fee", "label": "服务费用", "type": "number", "group": "金额信息", "priority": 4},
        ],
        "租赁合同": [
            {"key": "rental_item", "label": "租赁物", "type": "text", "group": "合同内容", "priority": 5},
            {"key": "rental_period", "label": "租赁期限", "type": "text", "group": "合同内容", "priority": 5},
            {"key": "rental_fee", "label": "租金", "type": "number", "group": "金额信息", "priority": 4},
            {"key": "rental_location", "label": "租赁物地址", "type": "text", "group": "合同内容", "priority": 5},
        ],
        "借款合同": [
            {"key": "loan_amount", "label": "借款金额", "type": "number", "group": "金额信息", "priority": 4},
            {"key": "loan_period", "label": "借款期限", "type": "text", "group": "合同内容", "priority": 5},
            {"key": "interest_rate", "label": "利率", "type": "text", "group": "金额信息", "priority": 4},
            {"key": "repayment_method", "label": "还款方式", "type": "text", "group": "合同内容", "priority": 5},
        ],
    }
    extras = category_extras.get(category, [])
    seen_keys = {f["key"] for f in base_fields}
    for ef in extras:
        if ef["key"] not in seen_keys:
            base_fields.append(ef)
            seen_keys.add(ef["key"])
    return sorted(base_fields, key=lambda f: f.get("priority", 99))


def _tpl_to_dict(tpl, detail=False):
    d = {
        "id": tpl.id, "name": tpl.name, "category": tpl.category,
        "description": tpl.description, "version": tpl.version,
        "status": tpl.status, "created_at": str(tpl.created_at),
        "updated_at": str(tpl.updated_at),
        "fields_count": len(tpl.fill_fields) if tpl.fill_fields else 0
    }
    if detail:
        d["fill_fields"] = tpl.fill_fields or []
        d["template_text"] = (tpl.template_text or "")[:2000]
    return d
